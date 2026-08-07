# -*- coding: utf-8 -*-
"""
Folha do EXERCICIO da aula 1.2 · pedidos-para-reescrever.docx

A folha e onde o aluno marca o que falta ANTES de reescrever. A ordem importa:
quem reescreve direto conserta so o que ja tinha percebido.

O que esta versao conserta, e o motivo:

  1. CADA CASO PASSOU A TER ARQUIVO. Antes a folha mandava rodar no Claude e
     contar mensagens, e nao existia arquivo nenhum para anexar. O aluno colava
     "Anexei tres propostas..." numa conversa vazia e o Claude respondia que nao
     tinha recebido anexo. A nota da aula era impossivel de medir.

  2. COMPRAS DEIXOU DE REPETIR A DEMONSTRACAO DA 1.1. O pedido antigo era
     "veja essas cotacoes e me diga qual e a melhor", que e literalmente a
     demonstracao da aula 1.1. Quem e de Compras ja tinha visto a resposta.
     Agora o pedido e outro sobre o mesmo arquivo: montar o pedido de compra.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TINTA = RGBColor(0x1F, 0x3A, 0x47)
CINZA = RGBColor(0x60, 0x6A, 0x70)
LINHA = "_" * 88

CASOS = [
    ("COMPRAS",
     "cotacoes-fornecedores.xlsx",
     "104 linhas · 26 insumos cotados por 4 fornecedores",
     '"Dá uma olhada nessas cotações e monta o pedido de compra."'),
    ("FINANCEIRO",
     "fechamento-dois-meses.xlsx",
     "118 linhas · fechamento de junho e de julho, uma aba por mês",
     '"Analisa esses números do mês e me diz o que aconteceu."'),
    ("RH E DP",
     "inscricoes-vaga-auxiliar-offset.xlsx",
     "124 inscrições · a descrição da vaga está na segunda aba",
     '"Vê quais desses candidatos servem para a vaga."'),
]

PERGUNTAS = [
    ('O que é "pronto"?', "Formato, tamanho, para quem vai ler"),
    ("Quais são as restrições?", "O que ignorar, o que não pode aparecer"),
    ("O que fazer na dúvida?", "Perguntar antes, ou seguir e declarar"),
]

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(1.9)
sec.left_margin = sec.right_margin = Cm(2.1)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)


def p(txt="", *, tam=10.5, negrito=False, cor=TINTA, antes=0, depois=4,
      italico=False, alinha=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(antes)
    par.paragraph_format.space_after = Pt(depois)
    if alinha:
        par.alignment = alinha
    r = par.add_run(txt)
    r.font.size = Pt(tam)
    r.font.bold = negrito
    r.font.italic = italico
    r.font.color.rgb = cor
    return par


def linhas_para_escrever(n):
    for _ in range(n):
        p(LINHA, tam=10, cor=RGBColor(0xA8, 0xB2, 0xB8), depois=8)


# ── cabeçalho ───────────────────────────────────────────────────────────────
p("Pedidos para reescrever", tam=20, negrito=True, depois=2)
p("Aula 1.2 · Pedir para entregar   ·   Claude na Prática   ·   Gráfica Aurora",
  tam=9.5, cor=CINZA, depois=12)

p("Escolha o caso do setor mais parecido com o seu. Baixe o arquivo indicado, "
  "marque com um X quais das três perguntas o pedido original NÃO responde, e "
  "só depois escreva a sua versão.", depois=4)
p("A ordem importa: marcar primeiro é o que faz você enxergar o buraco. Quem "
  "reescreve direto costuma consertar só o que já tinha percebido.",
  italico=True, cor=CINZA, depois=14)

# ── os três casos ───────────────────────────────────────────────────────────
for i, (setor, arquivo, sobre, pedido) in enumerate(CASOS, start=1):
    p(f"{i}.  {setor}", tam=13, negrito=True, antes=10, depois=3)
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(2)
    r = par.add_run("Arquivo:  ")
    r.font.size = Pt(10)
    r.font.color.rgb = CINZA
    r = par.add_run(arquivo)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = TINTA
    p(sobre, tam=9.5, italico=True, cor=CINZA, depois=6)

    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(8)
    r = par.add_run("O pedido como ele foi feito:  ")
    r.font.size = Pt(10)
    r.font.color.rgb = CINZA
    r = par.add_run(pedido)
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = TINTA

    tab = doc.add_table(rows=1, cols=2)
    tab.style = "Table Grid"
    tab.columns[0].width = Cm(1.1)
    cab = tab.rows[0].cells
    cab[0].text = "X"
    cab[1].text = "O que o pedido original NÃO responde"
    for cel in cab:
        for par in cel.paragraphs:
            for r in par.runs:
                r.font.size = Pt(9)
                r.font.bold = True
    for n, (titulo, dica) in enumerate(PERGUNTAS, start=1):
        cel = tab.add_row().cells
        cel[0].text = ""
        par = cel[1].paragraphs[0]
        r = par.add_run(f"{n}. {titulo}  ")
        r.font.size = Pt(10)
        r.font.bold = True
        r = par.add_run(dica)
        r.font.size = Pt(9.5)
        r.font.color.rgb = CINZA

    p("Reescreva aqui:", tam=10, negrito=True, antes=8, depois=6)
    linhas_para_escrever(5)
    p("Quantas idas e vindas até a resposta servir?  ________",
      tam=10, negrito=True, depois=14)

# ── o pedido de verdade ─────────────────────────────────────────────────────
doc.add_page_break()
p("Agora o de verdade", tam=15, negrito=True, depois=4)
p("Pegue um pedido que VOCÊ fez ao Claude nesta semana e que deu trabalho. Se "
  "não usou o Claude, use um pedido que você fez para uma pessoa: um e-mail "
  "pedindo relatório, um WhatsApp pedindo levantamento. Funciona igual, e a "
  "descoberta costuma ser mais desconfortável.", depois=12)

p("O pedido como você fez:", tam=10, negrito=True, depois=6)
linhas_para_escrever(3)
p("Quantas idas e vindas ele custou?  ________", tam=10, negrito=True, depois=14)

p("Reescrito, com as três perguntas respondidas:", tam=10, negrito=True, depois=6)
linhas_para_escrever(7)
p("Quantas idas e vindas agora?  ________", tam=10, negrito=True, depois=10)

p("Esse segundo número é o que você leva desta aula. Anote também onde ele "
  "parou de cair: a partir dali o problema deixou de ser o pedido.",
  tam=9.5, italico=True, cor=CINZA)

destino = ("/Users/rafaellima/developer/4-cursos-treinamentos/treinamentos-in-company/"
           "pouchain-claude-na-pratica/site/m1/a2-pedir-para-entregar/exercicio/"
           "pedidos-para-reescrever.docx")
doc.save(destino)

print(f"gravado em {destino}\n")
print(f"{'casos na folha':28} {len(CASOS)}")
for setor, arquivo, sobre, pedido in CASOS:
    print(f"{'':28} {setor:12} -> {arquivo}")
print(f"{'perguntas por caso':28} {len(PERGUNTAS)}")
print("\ntodo caso tem arquivo: o exercício agora roda de verdade e a")
print("contagem de idas e vindas passa a ser mensurável.")
